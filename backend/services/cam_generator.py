"""
CAM & Sanction Letter Document Generator
Uses python-docx for Word and reportlab for PDF generation.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Any, Dict, Optional
from datetime import datetime
import io

from services.supabase_client import get_supabase


def generate_cam_document(
    application_id: str,
    cam_content: Dict[str, Any],
    format: str = "docx",
) -> bytes:
    """Generate CAM document in Word or PDF format."""
    if format == "docx":
        return _generate_cam_docx(application_id, cam_content)
    elif format == "pdf":
        return _generate_cam_pdf(application_id, cam_content)
    else:
        raise ValueError(f"Unsupported format: {format}")


def _generate_cam_docx(application_id: str, cam_content: Dict[str, Any]) -> bytes:
    """Generate professional CAM Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("CREDIT APPRAISAL MEMO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Application ID: {application_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d-%b-%Y %H:%M')}")
    doc.add_paragraph(f"Confidential — For Internal Use Only")
    doc.add_paragraph("")

    # Sections from cam_content
    sections = cam_content.get("sections", {})
    if isinstance(cam_content, str):
        # If content is raw text from LLM
        doc.add_paragraph(cam_content)
    else:
        for section_title, section_body in sections.items():
            doc.add_heading(section_title, level=1)
            if isinstance(section_body, str):
                doc.add_paragraph(section_body)
            elif isinstance(section_body, dict):
                for key, value in section_body.items():
                    doc.add_paragraph(f"{key}: {value}", style="List Bullet")
            elif isinstance(section_body, list):
                for item in section_body:
                    doc.add_paragraph(str(item), style="List Bullet")

    # Footer
    doc.add_paragraph("")
    doc.add_paragraph("--- End of Credit Appraisal Memo ---").alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _generate_cam_pdf(application_id: str, cam_content: Dict[str, Any]) -> bytes:
    """Generate CAM PDF using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("CREDIT APPRAISAL MEMO", styles["Title"]))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Application ID: {application_id}", styles["Normal"]))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%d-%b-%Y %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 20))

    if isinstance(cam_content, str):
        story.append(Paragraph(cam_content, styles["Normal"]))
    else:
        sections = cam_content.get("sections", {})
        for title, body in sections.items():
            story.append(Paragraph(title, styles["Heading1"]))
            if isinstance(body, str):
                story.append(Paragraph(body, styles["Normal"]))
            story.append(Spacer(1, 10))

    doc.build(story)
    return buffer.getvalue()


def generate_sanction_letter_doc(application_id: str, format: str = "docx") -> Dict[str, Any]:
    """Generate professional sanction letter from DB data."""
    supabase = get_supabase()

    # Get application + decision data
    app = supabase.table("applications").select("*").eq("id", application_id).single().execute()
    decision = supabase.table("loan_decisions").select("*").eq(
        "application_id", application_id
    ).eq("action", "approve").order("created_at", desc=True).limit(1).execute()

    app_data = app.data
    dec_data = decision.data[0] if decision.data else {}

    doc = Document()
    doc.add_heading("SANCTION LETTER", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%b-%Y')}")
    doc.add_paragraph(f"To: {app_data.get('company_name', 'N/A')}")
    doc.add_paragraph("")
    doc.add_paragraph("Dear Sir/Madam,")
    doc.add_paragraph("")
    doc.add_paragraph(
        f"We are pleased to inform you that your loan application (Reference: {application_id}) "
        f"for {app_data.get('loan_type', 'N/A')} facility has been sanctioned with the following terms:"
    )

    # Terms table
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    terms = [
        ("Facility Type", app_data.get("loan_type", "N/A")),
        ("Sanctioned Limit", f"INR {dec_data.get('approved_limit', 'N/A')} Lakhs"),
        ("Interest Rate", f"{dec_data.get('approved_rate', 'N/A')}% p.a."),
        ("Tenure", f"{dec_data.get('approved_tenure_months', 'N/A')} months"),
        ("Risk Grade", app_data.get("final_risk_grade", "N/A")),
    ]
    for i, (label, value) in enumerate(terms):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)

    # Conditions
    conditions = dec_data.get("conditions", [])
    if conditions:
        doc.add_heading("Conditions Precedent", level=2)
        for c in conditions:
            doc.add_paragraph(c, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("For and on behalf of the Bank")
    doc.add_paragraph("Authorized Signatory")

    buffer = io.BytesIO()
    doc.save(buffer)
    file_bytes = buffer.getvalue()

    # Upload to Supabase Storage
    storage_path = f"sanction_letters/{application_id}/sanction_letter.docx"
    try:
        supabase.storage.from_("documents").upload(path=storage_path, file=file_bytes)
        file_url = supabase.storage.from_("documents").get_public_url(storage_path)
    except Exception:
        file_url = ""

    return {
        "application_id": application_id,
        "letter_url": file_url,
        "format": format,
        "generated_at": datetime.utcnow().isoformat(),
    }

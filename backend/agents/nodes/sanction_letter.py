"""
Agent 8: Sanction Letter Generator — Auto-generates professional sanction/rejection
letter using python-docx.

Triggered only after sanctioning authority approves/rejects.
Reads final decision from loan_decisions table.
Generates DOCX, uploads to Supabase Storage, updates state.
"""

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.state import CreditApplicationState
from services.supabase_client import get_supabase


# ---------------------------------------------------------------------------
# DB helpers
# ---------------------------------------------------------------------------

def _fetch_decision(application_id: str) -> Optional[Dict[str, Any]]:
    """Fetch latest decision from loan_decisions table."""
    try:
        supabase = get_supabase()
        result = supabase.table("loan_decisions").select("*").eq(
            "application_id", application_id
        ).order("created_at", desc=True).limit(1).execute()
        return result.data[0] if result.data else None
    except Exception as exc:
        print(f"[SanctionLetter] Failed to fetch decision: {exc}")
        return None


def _fetch_application(application_id: str) -> Dict[str, Any]:
    """Fetch application data."""
    try:
        supabase = get_supabase()
        result = supabase.table("applications").select("*").eq(
            "id", application_id
        ).single().execute()
        return result.data or {}
    except Exception as exc:
        print(f"[SanctionLetter] Failed to fetch app: {exc}")
        return {}


# ---------------------------------------------------------------------------
# Reference number generator
# ---------------------------------------------------------------------------

def _generate_ref_number(application_id: str) -> str:
    """Generate reference: SL-{app_id[:8]}-{year}."""
    year = datetime.utcnow().year
    prefix = application_id[:8].upper().replace("-", "")
    return f"SL-{prefix}-{year}"


# ---------------------------------------------------------------------------
# DOCX generation — Sanction Letter
# ---------------------------------------------------------------------------

def _build_sanction_letter(
    app_data: Dict[str, Any],
    decision: Dict[str, Any],
    application_id: str,
) -> bytes:
    """Build professional sanction letter DOCX."""
    doc = Document()

    # ── Style defaults ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ── Header: Bank letterhead placeholder ──
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("[BANK LETTERHEAD]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("")  # spacer

    # ── Date & Reference ──
    ref_number = _generate_ref_number(application_id)
    today = datetime.utcnow().strftime("%d %B %Y")

    doc.add_paragraph(f"Date: {today}")
    doc.add_paragraph(f"Ref: {ref_number}")
    doc.add_paragraph("")

    # ── Addressee ──
    company_name = app_data.get("company_name", "M/s. [Company Name]")
    address = app_data.get("registered_address", "[Registered Address]")
    doc.add_paragraph(f"To,")
    doc.add_paragraph(f"{company_name}")
    doc.add_paragraph(f"{address}")
    doc.add_paragraph("")

    # ── Subject ──
    loan_type = app_data.get("loan_type", "Credit Facility")
    sanctioned_amount = decision.get("approved_limit", app_data.get("loan_amount_requested", "N/A"))
    subject_para = doc.add_paragraph()
    run = subject_para.add_run(
        f"Subject: Sanction of {loan_type} facility of ₹{sanctioned_amount} Cr"
    )
    run.bold = True
    run.underline = True
    doc.add_paragraph("")

    # ── Salutation ──
    doc.add_paragraph("Dear Sir/Madam,")
    doc.add_paragraph("")

    # ── Opening ──
    doc.add_paragraph(
        f"With reference to your application for credit facility, we are pleased to inform you "
        f"that the bank has sanctioned the following facility in your favour, subject to the "
        f"terms and conditions mentioned herein:"
    )
    doc.add_paragraph("")

    # ── Terms table ──
    base_rate = decision.get("base_rate", 9.0)
    risk_premium = decision.get("risk_premium", decision.get("approved_rate", 2.0))
    total_rate = decision.get("approved_rate", float(base_rate) + float(risk_premium))
    tenure = decision.get("approved_tenure_months", 60)

    terms = [
        ("1. Sanctioned Amount", f"₹{sanctioned_amount} Cr"),
        ("2. Facility Type", loan_type),
        ("3. Purpose", app_data.get("purpose", "Business requirements")),
        ("4. Rate of Interest", f"{total_rate}% p.a."),
        ("5. Tenure", f"{tenure} months"),
        ("6. Repayment", f"Equal monthly instalments over {tenure} months"),
        ("7. Risk Grade", decision.get("risk_grade", app_data.get("final_risk_grade", "N/A"))),
    ]

    table = doc.add_table(rows=len(terms), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(terms):
        cell_l = table.cell(i, 0)
        cell_r = table.cell(i, 1)
        cell_l.text = label
        cell_r.text = str(value)
        for cell in (cell_l, cell_r):
            for para in cell.paragraphs:
                para.style.font.size = Pt(10)

    doc.add_paragraph("")

    # ── Collateral / Security ──
    collateral_data = decision.get("collateral_details") or app_data.get("collateral_details")
    doc.add_heading("Security / Collateral", level=2)
    if collateral_data:
        if isinstance(collateral_data, list):
            for item in collateral_data:
                doc.add_paragraph(str(item), style="List Bullet")
        elif isinstance(collateral_data, str):
            doc.add_paragraph(collateral_data)
        else:
            doc.add_paragraph(str(collateral_data))
    else:
        doc.add_paragraph("As per bank's standard norms for the facility type.")

    # ── Special Conditions ──
    conditions = decision.get("conditions") or []
    if conditions:
        doc.add_heading("Special Conditions", level=2)
        for c in conditions:
            doc.add_paragraph(str(c), style="List Bullet")

    # ── Validity ──
    doc.add_paragraph("")
    doc.add_paragraph(
        "Validity: This sanction is valid for a period of 90 days from the date of this letter. "
        "The borrower is required to comply with all conditions precedent within this period."
    )

    # ── Acceptance clause ──
    doc.add_paragraph("")
    doc.add_paragraph(
        "Acceptance: Please sign and return a copy of this letter as a token of your acceptance "
        "of the above terms and conditions. Disbursement will be subject to completion of all "
        "documentation and compliance with conditions precedent."
    )

    # ── Footer ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.add_run("For and on behalf of [Bank Name]").bold = True
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("________________________")
    doc.add_paragraph("Authorized Signatory")
    doc.add_paragraph("Name: [Sanctioning Authority]")
    doc.add_paragraph("Designation: [Designation]")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Upload + DB update
# ---------------------------------------------------------------------------

def _upload_letter(application_id: str, docx_bytes: bytes, letter_type: str) -> str:
    """Upload to Supabase Storage bucket."""
    try:
        supabase = get_supabase()
        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        path = f"sanction-letters/{application_id}/{letter_type}_{ts}.docx"
        supabase.storage.from_("documents").upload(path=path, file=docx_bytes)
        return supabase.storage.from_("documents").get_public_url(path)
    except Exception as exc:
        print(f"[SanctionLetter] Upload failed: {exc}")
        return ""


def _update_decision_record(application_id: str, letter_url: str) -> None:
    """Update loan_decisions with sanction letter URL."""
    try:
        supabase = get_supabase()
        supabase.table("loan_decisions").update({
            "sanction_letter_url": letter_url,
            "updated_at": datetime.utcnow().isoformat(),
        }).eq("application_id", application_id).execute()
    except Exception as exc:
        print(f"[SanctionLetter] DB update failed: {exc}")


# ---------------------------------------------------------------------------
# Main node
# ---------------------------------------------------------------------------

async def sanction_letter_node(
    state: CreditApplicationState,
) -> CreditApplicationState:
    """
    LangGraph node — Sanction Letter Generator.

    Triggered after sanctioning authority decision.
    Generates DOCX, uploads, updates DB and state.
    """
    state["current_node"] = "sanction_letter"
    state["progress_percent"] = 99
    state["status_message"] = "Generating sanction letter..."

    errors: list = list(state.get("errors") or [])
    application_id = state.get("application_id", "")

    if not application_id:
        errors.append("[sanction_letter] No application_id")
        state["errors"] = errors
        return state

    # ── Fetch decision ──
    decision = _fetch_decision(application_id)
    if not decision:
        state["status_message"] = "No decision found — sanction letter skipped."
        state["progress_percent"] = 100
        return state

    action = decision.get("action", "").lower()

    # ── Fetch application ──
    app_data = _fetch_application(application_id)

    if action == "approve":
        # ── Generate Sanction Letter ──
        try:
            docx_bytes = _build_sanction_letter(app_data, decision, application_id)
            letter_url = _upload_letter(application_id, docx_bytes, "sanction_letter")
            state["sanction_letter_url"] = letter_url
            _update_decision_record(application_id, letter_url)
            state["status_message"] = "Sanction letter generated and uploaded."
        except Exception as exc:
            errors.append(f"[sanction_letter] Generation failed: {exc}")
            state["status_message"] = f"Sanction letter failed: {exc}"

    elif action == "reject":
        state["status_message"] = "Application rejected — no sanction letter generated."
        state["rejection_letter_url"] = ""

    else:
        state["status_message"] = f"Decision action '{action}' — no letter generated."

    state["errors"] = errors
    state["progress_percent"] = 100
    state["completed_at"] = datetime.utcnow().isoformat()

    return state
